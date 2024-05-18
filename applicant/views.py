import random

from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, HttpResponseRedirect, FileResponse
from django.views.generic.edit import DeleteView

from django.urls import reverse

from .models import ApplicantPersonalFile, CompleteFrom, LanguageMath, ForeignLanguage, Privilege, Subject, VVK, \
    AddressRegion, AddressDistrict, AddressCity
from .filters import ApplicantFilter
from .forms import ApplicantForm
from faker import Faker
import csv
from io import BytesIO
import pandas as pd
from docx import Document


@login_required
def applicant_list(request):
    request.session['back_path'] = '/applicants/?' + request.META.get('QUERY_STRING')
    qs = ApplicantPersonalFile.objects.all()
    if 'o' in request.GET:
        order_query = request.GET['o'].split('.')
        qs = qs.order_by(*order_query)
    f = ApplicantFilter(request.GET, queryset=qs)
    paginator = Paginator(f.qs, 50)
    page = request.GET.get('page')
    applicant_list = paginator.get_page(page)
    applicant_form = ApplicantForm()
    return render(request, 'applicant/applicant_list.html',
                  {'applicant_list': applicant_list, 'applicant_form': applicant_form, 'filter': f})


@login_required
def add_applicant(request):
    if request.method == 'POST':
        form = ApplicantForm(request.POST)
        if form.is_valid():
            form.save()
            if 'from_modal' in request.POST:
                return HttpResponseRedirect(reverse('applicant:applicant-list'))
            else:
                return HttpResponseRedirect(reverse('applicant:applicant-add-form'))
        else:
            return render(request, 'applicant/add_form.html', {'applicant_form': form})
    if request.method == 'GET':
        applicant_form = ApplicantForm()
        return render(request, 'applicant/add_form.html', {'applicant_form': applicant_form})
    else:
        pass


@login_required
def update_applicant(request, applicant_id):
    if request.method == 'POST':
        obj = get_object_or_404(ApplicantPersonalFile, pk=applicant_id)
        form = ApplicantForm(request.POST, instance=obj)
        if form.is_valid():
            form.save()
            back_path = request.session.get('back_path', '/')
            return HttpResponseRedirect(back_path)
        else:
            return render(request, 'applicant/update_form.html', {'applicant_form': form,
                                                                  'obj': obj,
                                                                  })
    else:
        obj = get_object_or_404(ApplicantPersonalFile, pk=applicant_id)
        form = ApplicantForm(instance=obj)
        return render(request, 'applicant/update_form.html', {'applicant_form': form,
                                                              'obj': obj})


class ApplicantDelete(LoginRequiredMixin, DeleteView):
    model = ApplicantPersonalFile

    def get_success_url(self):
        return self.request.session.get('back_path', '/')


@login_required
def init_db(request):
    if request.method == 'POST':
        pass
    else:
        import csv
        AddressRegion.objects.all().delete()
        AddressDistrict.objects.all().delete()
        AddressCity.objects.all().delete()

        with open('by-list.csv') as csv_file:
            csv_reader = csv.reader(csv_file, delimiter=',')
            for row in csv_reader:
                region, created = AddressRegion.objects.get_or_create(region=row[0])
                district, created = AddressDistrict.objects.get_or_create(district=row[1], region=region)
                AddressCity.objects.get_or_create(city=row[2], district=district)

        CompleteFrom.objects.all().delete()

        with open('complete_from.csv') as csv_file:
            csv_reader = csv.reader(csv_file, delimiter=',')
            for row in csv_reader:
                CompleteFrom.objects.create(pk=row[0], complete_from=row[1])

        LanguageMath.objects.all().delete()
        LanguageMath.objects.create(pk=1, language="русский")
        LanguageMath.objects.create(pk=2, language="белоруский")

        ForeignLanguage.objects.all().delete()
        ForeignLanguage.objects.create(pk=1, foreign_language="анлийский")
        ForeignLanguage.objects.create(pk=2, foreign_language="немецкий")

        Privilege.objects.all().delete()
        Privilege.objects.create(pk=1, privilege='сирота')
        Privilege.objects.create(pk=2, privilege='опека')
        Subject.objects.all().delete()
        Subject.objects.create(pk=1, subject='русский язык')
        Subject.objects.create(pk=2, subject='белорусский язык')
        VVK.objects.all().delete()
        VVK.objects.create(pk=1, vvk_result="годен")
        VVK.objects.create(pk=2, vvk_result="не годен")

        fake = Faker('ru_RU')
        for i in range(10000):
            random_complete_from = CompleteFrom.objects.get(pk=random.randint(1, 2))
            random_privilege = Privilege.objects.get(pk=random.randint(1, 2))
            random_language_math = LanguageMath.objects.get(pk=random.randint(1, 2))
            random_subject = Subject.objects.get(pk=random.randint(1, 2))
            random_foreign_language = ForeignLanguage.objects.get(pk=random.randint(1, 2))
            random_vvk = VVK.objects.get(pk=random.randint(1, 2))
            ApplicantPersonalFile.objects.create(
                last_name=fake.last_name(),
                first_name=fake.first_name(),
                date_of_birth=fake.date_of_birth(minimum_age=14, maximum_age=16),
                address=fake.address(),
                contact_number=fake.phone_number(),
                complete_from=random_complete_from,
                average_mark=fake.pyfloat(right_digits=1, positive=True, min_value=6, max_value=9),
                class_he_goes_to=fake.pyint(min_value=7, max_value=10),
                there_is_application=fake.pybool(),
                there_is_birth_certificate=fake.pybool(),
                there_is_mark_sheet=fake.pybool(),
                there_is_characteristic=fake.pybool(),
                there_is_certificate_of_education=fake.pybool(),
                privilege=random_privilege,
                there_is_conclusion=fake.pybool(),
                there_is_medical_certificate=fake.pybool(),
                there_is_card_extract=fake.pybool(),
                there_is_psychological_information=fake.pybool(),
                language_math=random_language_math,
                language_for_dictation=random_subject,
                rus_mark=fake.pyint(min_value=7, max_value=10),
                bel_mark=fake.pyint(min_value=7, max_value=10),
                math_mark=fake.pyint(min_value=7, max_value=10),
                sum_mark=fake.pyint(min_value=7, max_value=10),
                language_for_study=random_foreign_language,
                fit=random_vvk
            )

        return HttpResponse("DB init is completed!!!")


def make_excel(queryset):
    file_name = 'Sample.xls'
    byte_buffer = BytesIO()
    df = pd.DataFrame.from_dict(
        queryset.values('last_name',
                        'first_name',
                        'patronymic',
                        'date_of_birth',
                        'address_city__city',
                        'address',
                        'contact_number',
                        'complete_from__complete_from',
                        'average_mark',
                        'average_mark_year',
                        'class_he_goes_to',
                        'there_is_application',
                        'there_is_birth_certificate',
                        'there_is_mark_sheet',
                        'there_is_characteristic',
                        'there_is_certificate_of_education',
                        'privilege',
                        'there_is_conclusion',
                        'there_is_medical_certificate',
                        'there_is_card_extract',
                        'there_is_psychological_information',
                        'language_math',
                        'language_for_dictation',
                        'rus_mark',
                        'bel_mark',
                        'math_mark',
                        'sum_mark',
                        'language_for_study',
                        'fit',
                        'fit_diagnosis',
                        'extra_data'
                        ))
    df.columns = ['Фамилия',
                  'Имя',
                  'Отчество',
                  'Дата рождения',
                  'Город',
                  'Адрес',
                  'Контактные данные',
                  'Комплектующий орган',
                  'Средний бал свидетельства об общем базовом образовании',
                  'Средний бал ведомости годовых отметок',
                  'В какой класс поступает',
                  'Заявление (да/нет)',
                  'Свидетельство о рождении (да/нет)',
                  'Ведомость годовых отметок (да/нет)',
                  'Характеристика (да/нет)',
                  'Свидетельство об общем базовом образовании (да/нет)',
                  'Документы, подтверждающие право на льготы',
                  'Заключение об изучении кандидата (да/нет)',
                  'Медицинская справка (да/нет)',
                  'Выписка из медицинской карты за последние 5 лет  (да/нет)',
                  'Медицинская справка о состоянии здоровья, подтверждающая отсутствие учета',
                  'Язык задания по математике',
                  'Выбор учебного предмета для написания диктанта',
                  'Отметка по Русский язык',
                  'Отметка по Белорусский язык',
                  'Отметка по Математика',
                  'Сумма баллов по результатам вступительных испытаний',
                  'Иностранный язык в случае поступления',
                  'Прохождение ВВК',
                  'Диагноз (если не годен)',
                  'Примечание'
                  ]
    df['Заявление (да/нет)'] = df['Заявление (да/нет)'].replace({True: 'Да', False: 'Нет'})
    df['Свидетельство о рождении (да/нет)'] = df['Свидетельство о рождении (да/нет)'].replace(
        {True: 'Да', False: 'Нет'})

    writer = pd.ExcelWriter(byte_buffer, engine='xlsxwriter')
    df.to_excel(writer, sheet_name='Employee', index=False)
    writer.close()
    return byte_buffer, file_name


def make_document(queryset):
    """generate docx response"""
    df = pd.DataFrame.from_dict(queryset[:10].values('last_name',
                                                     'first_name',
                                                     'patronymic',
                                                     'date_of_birth',
                                                     'address_city__city',
                                                     'address',
                                                     'contact_number',
                                                     'complete_from__complete_from',
                                                     'average_mark',
                                                     'average_mark_year',
                                                     'class_he_goes_to',
                                                     'there_is_application',
                                                     'there_is_birth_certificate',
                                                     'there_is_mark_sheet',
                                                     'there_is_characteristic',
                                                     'there_is_certificate_of_education',
                                                     'privilege',
                                                     'there_is_conclusion',
                                                     'there_is_medical_certificate',
                                                     'there_is_card_extract',
                                                     'there_is_psychological_information',
                                                     'language_math',
                                                     'language_for_dictation',
                                                     'rus_mark',
                                                     'bel_mark',
                                                     'math_mark',
                                                     'sum_mark',
                                                     'language_for_study',
                                                     'fit',
                                                     'fit_diagnosis',
                                                     'extra_data'
                                                     ))
    file_name = "Sample.docx"
    document = Document()
    document.add_heading("Employee List Sample")
    document.add_paragraph("This is sample document for employee list detail")
    table = document.add_table(rows=1, cols=len(df.columns.values.tolist()))
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns.values.tolist()):
        hdr_cells[i].text = col

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns.values.tolist()):
            row_cells[i].text = str(row[col])
    document.add_page_break()
    byte_buffer = BytesIO()
    document.save(byte_buffer)  # save your memory stream
    return byte_buffer, file_name


def get_excel(request):
    qs = ApplicantPersonalFile.objects.all()
    if 'o' in request.GET:
        order_query = request.GET['o'].split('.')
        qs = qs.order_by(*order_query)
    f = ApplicantFilter(request.GET, queryset=qs)
    filtered_qs = f.qs

    if 'res_type' in request.GET:
        res_type = request.GET['res_type']
        print('res_type', res_type)
    else:
        res_type = 'excel'

    if res_type == 'excel':
        byte_buffer, file_name = make_excel(filtered_qs)
        byte_buffer.seek(0)
    elif res_type == 'word':
        byte_buffer, file_name = make_document(filtered_qs)
        byte_buffer.seek(0)
    else:
        byte_buffer, file_name = make_excel(filtered_qs)
        byte_buffer.seek(0)
    return FileResponse(byte_buffer, filename=file_name, as_attachment=True)
